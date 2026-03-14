# DocxTableGenerator.py
# Generates .docx documents with compact tables for tabular document types.
# Used as an alternative to Google Apps Script for table-heavy documents.

import os
import tempfile
import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


TABULAR_TYPES = [
    "Master-Transcript-of-Marks",
    "Baccalaureate-Transcript-of-Marks-V1",
    "Baccalaureate-Transcript-of-Marks-V2",
]

DOWNLOAD_FOLDER = os.path.join(os.path.dirname(__file__), "..", "Downloads")


def _set_cell_shading(cell, color_hex):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, top=0, bottom=0, left=28, right=28):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcMar_existing = tcPr.find(qn("w:tcMar"))
    if tcMar_existing is not None:
        tcPr.remove(tcMar_existing)
    tcPr.append(tcMar)


def _set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{int(height_pt * 20)}" w:hRule="atLeast"/>'
    )
    existing = trPr.find(qn("w:trHeight"))
    if existing is not None:
        trPr.remove(existing)
    trPr.append(trHeight)


def _format_document_type(doc_type):
    return doc_type.replace("-", " ").title()


class DocxTableGenerator:
    def is_tabular(self, session):
        return session.get("Information Type") == "Tabular" and \
               session.get("Document Type") in TABULAR_TYPES

    def generate(self, session):
        if not os.path.exists(DOWNLOAD_FOLDER):
            os.makedirs(DOWNLOAD_FOLDER)

        doc = Document()

        # Page margins: narrower for more table space
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

        translation = session.get("Translation", {})
        text_fields = translation.get("Text", {})
        tables_data = translation.get("Tables", {})
        doc_type = session.get("Document Type", "")

        # Title
        title = doc.add_heading(_format_document_type(doc_type), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        doc.add_paragraph("")  # spacer

        # Text fields section
        field_order = [
            "Student Name", "Student National Code", "Institute",
            "Student Level", "Student Option", "Province",
            "City of issue", "Date of issue", "Year"
        ]

        for key in field_order:
            value = text_fields.get(key, "")
            if value:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                label_run = p.add_run(f"{key}: ")
                label_run.bold = True
                label_run.font.size = Pt(9)
                label_run.font.name = "Arial"
                value_run = p.add_run(str(value))
                value_run.font.size = Pt(9)
                value_run.font.name = "Arial"

        # Any remaining fields not in the order list
        for key, value in text_fields.items():
            if key not in field_order and value:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                label_run = p.add_run(f"{key}: ")
                label_run.bold = True
                label_run.font.size = Pt(9)
                label_run.font.name = "Arial"
                value_run = p.add_run(str(value))
                value_run.font.size = Pt(9)
                value_run.font.name = "Arial"

        doc.add_paragraph("")  # spacer before tables

        # Tables section
        if doc_type == "Master-Transcript-of-Marks":
            self._add_master_table(doc, tables_data)
        elif doc_type in [
            "Baccalaureate-Transcript-of-Marks-V1",
            "Baccalaureate-Transcript-of-Marks-V2"
        ]:
            self._add_baccalaureate_tables(doc, tables_data)

        # Footer info
        doc.add_paragraph("")
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer_p.paragraph_format.space_before = Pt(4)
        sid_run = footer_p.add_run(f"Session ID: {session.get('Session Id', '')}")
        sid_run.font.size = Pt(7)
        sid_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        footer_p.add_run("    ")
        date_run = footer_p.add_run(f"Date: {session.get('Operation Date', '')}")
        date_run.font.size = Pt(7)
        date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Save the document
        session_id = session.get("Session Id", "document")
        filename = f"{session_id}.docx"
        filepath = os.path.join(DOWNLOAD_FOLDER, filename)
        doc.save(filepath)

        return filename

    def _add_compact_table(self, doc, headers, rows, title=None):
        if title:
            heading = doc.add_heading(title, level=3)
            heading.paragraph_format.space_before = Pt(4)
            heading.paragraph_format.space_after = Pt(2)
            for run in heading.runs:
                run.font.size = Pt(10)

        num_cols = len(headers)
        num_rows = len(rows) + 1  # +1 for header

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Calculate column widths based on available page width (~18cm with 1.5cm margins)
        available_width_cm = 18.0
        if num_cols <= 4:
            first_col_cm = available_width_cm * 0.40
            other_col_cm = (available_width_cm - first_col_cm) / (num_cols - 1)
        else:
            first_col_cm = available_width_cm * 0.18
            other_col_cm = (available_width_cm - first_col_cm) / (num_cols - 1)

        # Set column widths
        for col_idx in range(num_cols):
            width = Cm(first_col_cm if col_idx == 0 else other_col_cm)
            table.columns[col_idx].width = width

        # Header row
        header_row = table.rows[0]
        _set_row_height(header_row, 12)
        for col_idx, header_text in enumerate(headers):
            cell = header_row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(8)
            run = p.add_run(str(header_text))
            run.bold = True
            run.font.size = Pt(6.5)
            run.font.name = "Arial Narrow"
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            _set_cell_shading(cell, "D9D9D9")
            _set_cell_margins(cell, top=10, bottom=10, left=20, right=20)
            cell.width = Cm(first_col_cm if col_idx == 0 else other_col_cm)

        # Data rows
        for row_idx, row_data in enumerate(rows):
            table_row = table.rows[row_idx + 1]
            _set_row_height(table_row, 10)
            for col_idx in range(num_cols):
                cell = table_row.cells[col_idx]
                cell_value = ""
                if col_idx < len(row_data):
                    cell_value = str(row_data[col_idx]) if row_data[col_idx] is not None else ""
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(7.5)
                run = p.add_run(cell_value)
                run.font.size = Pt(6.5)
                run.font.name = "Arial Narrow"
                _set_cell_margins(cell, top=8, bottom=8, left=20, right=20)
                cell.width = Cm(first_col_cm if col_idx == 0 else other_col_cm)

                # Alternate row shading
                if row_idx % 2 == 1:
                    _set_cell_shading(cell, "F5F5F5")

        # Set table borders to thin
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BBBBBB"/>'
            f'</w:tblBorders>'
        )
        existing_borders = tblPr.find(qn("w:tblBorders"))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(borders)

    def _add_master_table(self, doc, tables_data):
        if not tables_data or not isinstance(tables_data, list):
            return

        headers = ["Subject", "Mark", "Result", "Session"]
        rows = []
        for item in tables_data:
            rows.append([
                item.get("Subject", ""),
                item.get("Mark", ""),
                item.get("Result", ""),
                item.get("Session", ""),
            ])

        self._add_compact_table(doc, headers, rows, title="Transcript of Marks")

    def _add_baccalaureate_tables(self, doc, tables_data):
        if not tables_data:
            return

        # Transcript table
        transcript = tables_data.get("Transcript")
        if transcript:
            columns = transcript.get("Columns", [])
            rows = transcript.get("Rows", [])
            if columns and rows:
                self._add_compact_table(doc, columns, rows, title="Transcript of Marks")

        # Overall table
        overall = tables_data.get("Overall")
        if overall:
            columns = overall.get("Columns", [])
            rows = overall.get("Rows", [])
            if columns and rows:
                self._add_compact_table(doc, columns, rows, title="Overall Results")
